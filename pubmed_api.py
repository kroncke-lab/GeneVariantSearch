import os
import time
from Bio import Entrez
import xml.etree.ElementTree as ET
from typing import List, Dict, Optional
import requests
import tempfile
import re
try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
try:
    from openpyxl import load_workbook
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

Entrez.email = os.environ.get("PUBMED_EMAIL", "user@example.com")

class PubMedAPI:
    def __init__(self, email: str = None):
        if email:
            Entrez.email = email
        self.max_results = 20
    
    def filter_relevant_sections(self, text: str, gene: str, keywords: Optional[List[str]] = None) -> str:
        """Extract paragraphs/sections that mention the gene or variant keywords"""
        if not text:
            return ""
        
        search_terms = [gene.lower()]
        if keywords:
            search_terms.extend([k.lower() for k in keywords])
        
        search_terms.extend(['variant', 'mutation', 'polymorphism', 'genotype', 'phenotype', 
                            'clinical', 'patient', 'carrier', 'heterozygous', 'homozygous'])
        
        paragraphs = re.split(r'\n\n+', text)
        relevant = []
        
        for para in paragraphs:
            para_lower = para.lower()
            if any(term in para_lower for term in search_terms):
                relevant.append(para)
        
        return '\n\n'.join(relevant)
    
    def parse_pdf(self, file_path: str, gene: str) -> str:
        """Extract text from PDF and filter for relevant content"""
        if not HAS_PDF:
            return ""
        
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                    
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            table_text = '\n'.join([' | '.join([str(cell) if cell else '' for cell in row]) for row in table])
                            text_parts.append(f"TABLE:\n{table_text}")
                
                full_text = '\n\n'.join(text_parts)
                return self.filter_relevant_sections(full_text, gene)
        except Exception as e:
            return f"[PDF parse error: {str(e)}]"
    
    def parse_excel(self, file_path: str, gene: str) -> str:
        """Extract text from Excel and filter for relevant content"""
        if not HAS_EXCEL:
            return ""
        
        try:
            wb = load_workbook(file_path, data_only=True)
            text_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"SHEET: {sheet_name}")
                
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                    if row_text.strip():
                        rows.append(row_text)
                
                text_parts.append('\n'.join(rows))
            
            full_text = '\n\n'.join(text_parts)
            return self.filter_relevant_sections(full_text, gene)
        except Exception as e:
            return f"[Excel parse error: {str(e)}]"
    
    def parse_docx(self, file_path: str, gene: str) -> str:
        """Extract text from Word document and filter for relevant content"""
        if not HAS_DOCX:
            return ""
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = ' | '.join([cell.text for cell in row.cells])
                    table_text.append(row_text)
                text_parts.append(f"TABLE:\n" + '\n'.join(table_text))
            
            full_text = '\n\n'.join(text_parts)
            return self.filter_relevant_sections(full_text, gene)
        except Exception as e:
            return f"[Word parse error: {str(e)}]"
    
    def download_and_parse_supplement(self, url: str, gene: str) -> Optional[str]:
        """Download a supplement file and parse based on file type"""
        try:
            response = requests.get(url, timeout=30)
            if response.status_code != 200:
                return None
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(url)[1]) as tmp:
                tmp.write(response.content)
                tmp_path = tmp.name
            
            try:
                file_ext = os.path.splitext(url)[1].lower()
                
                if file_ext == '.pdf':
                    return self.parse_pdf(tmp_path, gene)
                elif file_ext in ['.xlsx', '.xls']:
                    return self.parse_excel(tmp_path, gene)
                elif file_ext in ['.docx', '.doc']:
                    return self.parse_docx(tmp_path, gene)
                else:
                    return None
            finally:
                try:
                    os.unlink(tmp_path)
                except:
                    pass
        except Exception as e:
            return None
    
    def search_publications(self, gene: str, variant: Optional[str] = None, max_results: int = 20) -> List[str]:
        query = gene
        if variant:
            query = f"{gene} AND {variant}"
        
        query += " AND (genetic variant OR mutation OR polymorphism OR genotype)"
        
        try:
            search_handle = Entrez.esearch(
                db="pubmed",
                term=query,
                retmax=max_results,
                sort="relevance"
            )
            search_results = Entrez.read(search_handle)
            search_handle.close()
            
            id_list = search_results.get("IdList", [])
            return id_list
        except Exception as e:
            raise Exception(f"PubMed search failed: {str(e)}")
    
    def get_pmc_id(self, pmid: str) -> Optional[str]:
        """Get PMC ID for a given PMID if available in PubMed Central"""
        try:
            link_handle = Entrez.elink(
                dbfrom="pubmed",
                db="pmc",
                id=pmid,
                linkname="pubmed_pmc"
            )
            link_results = Entrez.read(link_handle)
            link_handle.close()
            
            if link_results and link_results[0].get("LinkSetDb"):
                pmc_ids = link_results[0]["LinkSetDb"][0]["Link"]
                if pmc_ids:
                    return pmc_ids[0]["Id"]
            return None
        except:
            return None
    
    def fetch_pmc_fulltext(self, pmc_id: str, gene: str) -> Optional[str]:
        """Fetch full text from PMC including supplemental data with intelligent filtering"""
        try:
            fetch_handle = Entrez.efetch(
                db="pmc",
                id=pmc_id,
                rettype="xml",
                retmode="xml"
            )
            xml_data = fetch_handle.read()
            fetch_handle.close()
            
            root = ET.fromstring(xml_data)
            
            full_text_parts = []
            
            title = root.find(".//article-title")
            if title is not None and title.text:
                full_text_parts.append(f"TITLE: {title.text}")
            
            abstract = root.find(".//abstract")
            if abstract is not None:
                abstract_text = " ".join([elem.text or "" for elem in abstract.iter() if elem.text])
                if abstract_text.strip():
                    full_text_parts.append(f"ABSTRACT: {abstract_text}")
            
            body = root.find(".//body")
            if body is not None:
                body_text = " ".join([elem.text or "" for elem in body.iter() if elem.text and elem.tag in ['p', 'title', 'td', 'th']])
                if body_text.strip():
                    filtered_body = self.filter_relevant_sections(body_text, gene)
                    if filtered_body:
                        full_text_parts.append(f"MAIN TEXT: {filtered_body[:30000]}")
            
            tables = root.findall(".//table-wrap")
            if tables:
                table_texts = []
                for idx, table in enumerate(tables[:15]):
                    table_content = " ".join([elem.text or "" for elem in table.iter() if elem.text])
                    if table_content.strip():
                        table_texts.append(f"Table {idx+1}: {table_content}")
                if table_texts:
                    full_text_parts.append(f"INLINE TABLES: {' | '.join(table_texts)}")
            
            supplementary = root.findall(".//supplementary-material")
            if supplementary:
                supp_count = 0
                for supp in supplementary:
                    supp_inline_text = " ".join([elem.text or "" for elem in supp.iter() if elem.text and elem.tag in ['caption', 'p', 'title']])
                    if supp_inline_text.strip():
                        full_text_parts.append(f"SUPPLEMENT CAPTION: {supp_inline_text}")
                    
                    media_links = supp.findall(".//media")
                    for media in media_links:
                        href = media.get("{http://www.w3.org/1999/xlink}href")
                        if href:
                            supp_url = f"https://www.ncbi.nlm.nih.gov/pmc/articles/PMC{pmc_id}/bin/{href}"
                            
                            if supp_count < 5:
                                parsed_content = self.download_and_parse_supplement(supp_url, gene)
                                if parsed_content:
                                    full_text_parts.append(f"SUPPLEMENT FILE ({href}): {parsed_content[:10000]}")
                                    supp_count += 1
                                time.sleep(0.5)
            
            if full_text_parts:
                return "\n\n".join(full_text_parts)
            return None
            
        except Exception as e:
            return None
    
    def fetch_article_details(self, pmid_list: List[str], gene: str) -> List[Dict]:
        if not pmid_list:
            return []
        
        articles = []
        
        try:
            fetch_handle = Entrez.efetch(
                db="pubmed",
                id=pmid_list,
                rettype="xml",
                retmode="xml"
            )
            
            xml_data = fetch_handle.read()
            fetch_handle.close()
            
            root = ET.fromstring(xml_data)
            
            for article in root.findall(".//PubmedArticle"):
                try:
                    pmid_elem = article.find(".//PMID")
                    pmid = pmid_elem.text if pmid_elem is not None else "Unknown"
                    
                    title_elem = article.find(".//ArticleTitle")
                    title = title_elem.text if title_elem is not None else "No title"
                    
                    abstract_texts = article.findall(".//AbstractText")
                    abstract = " ".join([
                        elem.text for elem in abstract_texts if elem.text
                    ]) if abstract_texts else "No abstract available"
                    
                    year_elem = article.find(".//PubDate/Year")
                    year = year_elem.text if year_elem is not None else "Unknown"
                    
                    authors = []
                    for author in article.findall(".//Author"):
                        lastname = author.find("LastName")
                        forename = author.find("ForeName")
                        if lastname is not None and forename is not None:
                            authors.append(f"{forename.text} {lastname.text}")
                    
                    author_str = ", ".join(authors[:3]) if authors else "Unknown authors"
                    if len(authors) > 3:
                        author_str += " et al."
                    
                    pmc_id = self.get_pmc_id(pmid)
                    content_type = "abstract"
                    full_text = f"{title}\n\n{abstract}"
                    
                    if pmc_id:
                        time.sleep(0.34)
                        pmc_fulltext = self.fetch_pmc_fulltext(pmc_id, gene)
                        if pmc_fulltext:
                            full_text = pmc_fulltext
                            content_type = "full_text_with_supplements"
                    
                    articles.append({
                        "pmid": pmid,
                        "title": title,
                        "abstract": abstract,
                        "year": year,
                        "authors": author_str,
                        "full_text": full_text,
                        "content_type": content_type,
                        "pmc_id": pmc_id
                    })
                    
                except Exception as e:
                    continue
            
            return articles
            
        except Exception as e:
            raise Exception(f"Failed to fetch article details: {str(e)}")
    
    def search_and_fetch(self, gene: str, variant: Optional[str] = None, max_results: int = 20) -> List[Dict]:
        pmid_list = self.search_publications(gene, variant, max_results)
        
        if not pmid_list:
            return []
        
        time.sleep(0.34)
        
        articles = self.fetch_article_details(pmid_list, gene)
        return articles
